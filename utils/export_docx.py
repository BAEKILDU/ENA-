"""경쟁력 분석 결과 Word 문서 생성 (pro2 템플릿)."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from data.mock_data import get_revenue_ideas


def build_analysis_docx(result: dict) -> bytes:
    """분석 결과 dict를 Word(.docx) 바이트로 변환."""
    overview = result.get("overview") or {}
    swot = result.get("swot") or {}
    summary = result.get("summary") or {}
    kpi = result.get("kpi") or {}
    score_details = result.get("score_details") or {}

    title = overview.get("title", result.get("title", "신규 콘텐츠"))
    genre = overview.get("genre", result.get("genre", "미정"))
    channel = overview.get("channel", "ENA")
    slot = overview.get("slot", result.get("slot", "미정"))
    cast_list = overview.get("cast", result.get("cast", []))
    cast_str = ", ".join(cast_list) if cast_list else "미정"
    logline = overview.get("logline", "미정")
    overall = summary.get("overall", result.get("overall", 0))
    source_label = "업로드 기획안" if result.get("source") == "upload" else "직접 작성"
    source_file = result.get("source_file", "")
    source_info = f"{source_label} · {source_file}" if source_file else source_label

    doc = Document()
    doc.add_heading("ENA 가치+ 경쟁력 분석 결과", level=0)
    doc.add_paragraph("신규 콘텐츠 가치+ · 분석 상세 리포트")
    doc.add_paragraph(f"출처: {source_info}")

    doc.add_heading("1. 기획안 핵심 요약", level=1)
    doc.add_paragraph(summary.get("intent", swot.get("intent_summary", "-")))
    doc.add_paragraph(f"총평: {summary.get('one_liner', swot.get('one_liner', '-'))}")
    doc.add_paragraph(f"종합 경쟁력 점수: {overall}/10")
    doc.add_paragraph(f"핵심 강점: {summary.get('key_strength', '-')}")
    doc.add_paragraph(f"치명적 약점/리스크: {summary.get('key_risk', '-')}")

    doc.add_heading("2. 콘텐츠 개요", level=1)
    doc.add_paragraph(f"콘텐츠명: {title}")
    doc.add_paragraph(f"장르: {genre}")
    doc.add_paragraph(f"편성채널: {channel}")
    doc.add_paragraph(f"편성시간: {slot}")
    doc.add_paragraph(f"주요출연자: {cast_str}")
    doc.add_paragraph(f"로그라인: {logline}")

    doc.add_heading("3. 기획안 분석 — 강점 및 약점", level=1)
    doc.add_heading("강점", level=2)
    for item in swot.get("strengths") or ["-"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("약점", level=2)
    for item in swot.get("weaknesses") or ["-"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. 분석 결론 제안", level=1)
    doc.add_heading("긍정 의견", level=2)
    for item in swot.get("positive", []):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("중립 의견", level=2)
    for item in swot.get("neutral", []):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("부정 의견", level=2)
    for item in swot.get("negative", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. 주요 데이터 지표", level=1)
    doc.add_paragraph(f"종합 경쟁력 지수: {kpi.get('overall', overall)}")
    doc.add_paragraph(f"필요 핵심 출연진 수: {kpi.get('required_cast', len(cast_list))}명")
    doc.add_paragraph(f"최적 편성 시간 제안: {kpi.get('best_slot', slot)}")
    doc.add_paragraph(f"동시간대 또는 유사 장르 경쟁 프로그램 수: {kpi.get('competitor_count', 0)}편")

    doc.add_heading("6. 10점 만점 경쟁력 세부 지표", level=1)
    for name, detail in score_details.items():
        doc.add_paragraph(
            f"{name}: {detail.get('score', '-')}/10 — {detail.get('reason', '-')}",
            style="List Bullet",
        )

    doc.add_heading("7. 부가 사업 및 수익 창출 아이디어", level=1)
    for group in get_revenue_ideas(genre):
        doc.add_heading(group["category"], level=2)
        for idea in group["ideas"]:
            doc.add_paragraph(idea, style="List Bullet")

    doc.add_heading("8. 종합 결론", level=1)
    doc.add_paragraph(
        result.get("final_conclusion")
        or swot.get("final_conclusion")
        or "종합 결론을 생성하려면 분석을 다시 실행해 주세요."
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
